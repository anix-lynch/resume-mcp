#!/usr/bin/env python3
"""
Extract text from CV 2019 docx file
"""

try:
    from docx import Document
    doc = Document("Anix Lynch CV 2019 for paulina.docx")
    
    print("=== CV 2019 Content ===\n")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n=== Tables ===\n")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                print(" | ".join(row_text))
                
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run(["pip3", "install", "--break-system-packages", "-q", "python-docx"])
    from docx import Document
    doc = Document("Anix Lynch CV 2019 for paulina.docx")
    
    print("=== CV 2019 Content ===\n")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n=== Tables ===\n")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                print(" | ".join(row_text))

